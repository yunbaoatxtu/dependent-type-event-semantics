#!/usr/bin/env python3
"""Regenerate the manuscript DOCX from the Markdown source."""

from __future__ import annotations

import argparse
from dataclasses import dataclass
from pathlib import Path
from typing import Any

try:
    from scripts.paper_markdown import is_table_separator, markdown_inline_segments, split_table_row
except ModuleNotFoundError:  # pragma: no cover - used when run as scripts/sync_paper_docx.py.
    from paper_markdown import is_table_separator, markdown_inline_segments, split_table_row


ROOT = Path(__file__).resolve().parents[1]
DEFAULT_MARKDOWN = ROOT / "paper" / "dependent_type_replacement_for_event_semantics_sci_manuscript.md"
DEFAULT_DOCX = ROOT / "paper" / "dependent_type_replacement_for_event_semantics_sci_manuscript.docx"


@dataclass(frozen=True)
class DocxTooling:
    Document: Any
    WD_ALIGN_PARAGRAPH: Any
    Inches: Any
    Pt: Any
    RGBColor: Any


def load_docx_tooling() -> DocxTooling:
    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Inches, Pt, RGBColor
    except ImportError as exc:  # pragma: no cover - depends on local document tooling.
        raise SystemExit(
            "python-docx is required. Run this script with the bundled Codex "
            "workspace Python runtime or install python-docx."
        ) from exc
    return DocxTooling(Document, WD_ALIGN_PARAGRAPH, Inches, Pt, RGBColor)


def add_markdown_runs(
    paragraph: Any,
    text: str,
    *,
    base_bold: bool = False,
    italic: bool = False,
    color: Any | None = None,
) -> None:
    for segment in markdown_inline_segments(text):
        run = paragraph.add_run(segment.text)
        if base_bold or segment.bold:
            run.bold = True
        if italic or segment.italic:
            run.italic = True
        if segment.code:
            run.font.name = "Courier New"
        if color is not None:
            run.font.color.rgb = color


def apply_base_styles(document: Any, tooling: DocxTooling) -> None:
    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = tooling.Pt(10.5)
    normal.paragraph_format.space_after = tooling.Pt(6)
    normal.paragraph_format.line_spacing = 1.12

    for name, size in [("Heading 1", 15), ("Heading 2", 12.5)]:
        style = styles[name]
        style.font.name = "Arial"
        style.font.size = tooling.Pt(size)
        style.font.bold = True
        style.font.color.rgb = tooling.RGBColor(43, 110, 171)
        style.paragraph_format.space_before = tooling.Pt(12)
        style.paragraph_format.space_after = tooling.Pt(6)

    title = styles["Title"]
    title.font.name = "Arial"
    title.font.size = tooling.Pt(16)
    title.font.bold = True
    title.font.color.rgb = tooling.RGBColor(31, 78, 121)
    title.paragraph_format.alignment = tooling.WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = tooling.Pt(4)


def add_keywords(document: Any, line: str) -> None:
    paragraph = document.add_paragraph()
    add_markdown_runs(paragraph, line)


def add_markdown_table(document: Any, rows: list[list[str]], tooling: DocxTooling) -> None:
    if not rows:
        return
    table = document.add_table(rows=1, cols=len(rows[0]))
    table.style = "Table Grid"
    table.autofit = True
    header_cells = table.rows[0].cells
    for index, value in enumerate(rows[0]):
        add_markdown_runs(header_cells[index].paragraphs[0], value, base_bold=True)

    for row in rows[1:]:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            add_markdown_runs(cells[index].paragraphs[0], value)

    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = tooling.Pt(2)


def build_docx(markdown_path: Path, docx_path: Path) -> None:
    tooling = load_docx_tooling()
    document = tooling.Document()
    section = document.sections[0]
    section.top_margin = tooling.Inches(0.9)
    section.bottom_margin = tooling.Inches(0.75)
    section.left_margin = tooling.Inches(0.9)
    section.right_margin = tooling.Inches(0.9)
    apply_base_styles(document, tooling)

    lines = markdown_path.read_text(encoding="utf-8").splitlines()
    index = 0
    while index < len(lines):
        line = lines[index].strip()
        if not line:
            index += 1
            continue

        if line.startswith("# "):
            paragraph = document.add_paragraph(style="Title")
            add_markdown_runs(paragraph, line[2:].strip(), base_bold=True)
            index += 1
            continue

        if line.startswith("_") and line.endswith("_") and len(line) > 1:
            paragraph = document.add_paragraph()
            paragraph.alignment = tooling.WD_ALIGN_PARAGRAPH.CENTER
            add_markdown_runs(
                paragraph,
                line.strip("_"),
                italic=True,
                color=tooling.RGBColor(91, 107, 120),
            )
            index += 1
            continue

        if line.startswith("## "):
            heading = line[3:].strip()
            if heading == "Functional replacement map":
                document.add_page_break()
            paragraph = document.add_paragraph(style="Heading 1")
            add_markdown_runs(paragraph, heading, base_bold=True)
            index += 1
            continue

        if line.startswith("|"):
            table_rows: list[list[str]] = []
            while index < len(lines) and lines[index].strip().startswith("|"):
                current = lines[index].strip()
                if not is_table_separator(current):
                    table_rows.append(split_table_row(current))
                index += 1
            add_markdown_table(document, table_rows, tooling)
            continue

        if line.startswith("- "):
            paragraph = document.add_paragraph(style="List Bullet")
            add_markdown_runs(paragraph, line[2:].strip())
            index += 1
            continue

        if line.startswith("**Keywords:**"):
            add_keywords(document, line)
            index += 1
            continue

        paragraph = document.add_paragraph()
        add_markdown_runs(paragraph, line)
        index += 1

    footer = document.sections[0].footer.paragraphs[0]
    footer.text = "Manuscript draft"
    footer.alignment = tooling.WD_ALIGN_PARAGRAPH.RIGHT
    docx_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(docx_path)


def main() -> None:
    parser = argparse.ArgumentParser(description="Synchronize the paper DOCX from Markdown.")
    parser.add_argument("--markdown", type=Path, default=DEFAULT_MARKDOWN)
    parser.add_argument("--docx", type=Path, default=DEFAULT_DOCX)
    args = parser.parse_args()
    build_docx(args.markdown, args.docx)
    print(f"wrote {args.docx}")


if __name__ == "__main__":
    main()
